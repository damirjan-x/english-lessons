"""
Генерация PDF-отчёта на основе шаблона .docx с подстановкой данных из БД.
В контейнере используется LibreOffice для конвертации docx -> pdf.
Подстановка плейсхолдеров [[ имя ]] — простая замена без Jinja2, чтобы символы
«<», «{» и т.п. в документе Word не вызывали ошибки парсера.
"""

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from app.config import get_settings

# Плейсхолдеры в шаблоне: [[ имя_переменной ]]
PLACEHOLDER_PATTERN = re.compile(r"\[\[\s*(\w+)\s*\]\]")

# Названия месяцев на русском; индекс 1–12 = январь–декабрь
MONTHS_RU = [
    "",  # 0 — не используется
    "январь", "февраль", "март", "апрель", "май", "июнь",
    "июль", "август", "сентябрь", "октябрь", "ноябрь", "декабрь",
]

# Международные (григорианские) названия месяцев на иврите; индекс 1–12 = январь–декабрь
MONTHS_HE = [
    "",  # 0 — не используется
    "ינואר",   # 1 январь
    "פברואר",  # 2 февраль
    "מרץ",     # 3 март
    "אפריל",   # 4 апрель
    "מאי",     # 5 май
    "יוני",    # 6 июнь
    "יולי",    # 7 июль
    "אוגוסט",  # 8 август
    "ספטמבר",  # 9 сентябрь
    "אוקטובר", # 10 октябрь
    "נובמבר",  # 11 ноябрь
    "דצמבר",   # 12 декабрь
]


def get_report_context(record) -> dict:
    """Собрать контекст для подстановки в шаблон (имена полей как в ТЗ). Все значения — строки."""
    student = record.student
    return {
        "first_name": str(student.first_name or ""),
        "last_name": str(student.last_name or ""),
        "first_name_he": str(student.first_name_he or ""),
        "last_name_he": str(student.last_name_he or ""),
        "year": str(record.year),
        "month": str(record.month),
        "month_ru": MONTHS_RU[record.month] if 1 <= record.month <= 12 else "",
        "month_he": MONTHS_HE[record.month] if 1 <= record.month <= 12 else "",
        "GrammarE": str(record.grammar_e or ""),
        "ReadingE": str(record.reading_e or ""),
        "SpeakingE": str(record.speaking_e or ""),
        "WritingE": str(record.writing_e or ""),
        "hours_studied": str(record.hours_studied if record.hours_studied is not None else 0),
    }


def _replace_placeholders(text: str, context: dict) -> str:
    """Заменить в тексте все [[ key ]] на context.get(key, '')."""
    if not text:
        return text
    def repl(match: re.Match) -> str:
        key = match.group(1)
        value = context.get(key, "")
        return str(value) if value is not None else ""
    return PLACEHOLDER_PATTERN.sub(repl, text)


def _render_docx_with_context(doc: Document, context: dict) -> None:
    """Подставить context во все параграфы и ячейки таблиц (изменяет doc)."""
    for para in doc.paragraphs:
        if PLACEHOLDER_PATTERN.search(para.text):
            para.text = _replace_placeholders(para.text, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if PLACEHOLDER_PATTERN.search(para.text):
                        para.text = _replace_placeholders(para.text, context)
    for section in doc.sections:
        try:
            if section.header is not None:
                for para in section.header.paragraphs:
                    if PLACEHOLDER_PATTERN.search(para.text):
                        para.text = _replace_placeholders(para.text, context)
        except Exception:
            pass
        try:
            if section.footer is not None:
                for para in section.footer.paragraphs:
                    if PLACEHOLDER_PATTERN.search(para.text):
                        para.text = _replace_placeholders(para.text, context)
        except Exception:
            pass


def render_docx_and_convert_to_pdf(record) -> bytes:
    """
    Заполнить шаблон .docx данными записи и сконвертировать в PDF.
    Возвращает содержимое PDF-файла.
    """
    settings = get_settings()
    base_dir = Path(__file__).resolve().parent.parent.parent
    templates_dir = base_dir / settings.templates_dir
    template_path = templates_dir / settings.report_template_name

    if not template_path.exists():
        raise FileNotFoundError(f"Report template not found: {template_path}")

    doc = Document(str(template_path))
    context = get_report_context(record)
    _render_docx_with_context(doc, context)

    tmpdir_parent = None
    if settings.temp_dir:
        p = Path(settings.temp_dir)
        if p.exists() and p.is_dir():
            tmpdir_parent = str(p)
    with tempfile.TemporaryDirectory(prefix="report_", dir=tmpdir_parent) as tmpdir:
        tmpdir_path = Path(tmpdir)
        docx_path = tmpdir_path / "report.docx"
        doc.save(str(docx_path))  # python-docx Document.save()

        # Конвертация в PDF через LibreOffice (в Docker) или локально
        pdf_path = _convert_docx_to_pdf(docx_path, tmpdir_path)

        if not pdf_path or not pdf_path.exists():
            raise RuntimeError("PDF conversion failed")

        return pdf_path.read_bytes()


def _convert_docx_to_pdf(docx_path: Path, out_dir: Path) -> Path | None:
    """
    Конвертация .docx в .pdf с помощью LibreOffice (headless).
    В контейнере должен быть установлен: apt-get install -y libreoffice-writer
    """
    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )
        pdf_path = out_dir / "report.pdf"
        return pdf_path if pdf_path.exists() else None
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        # Fallback: если LibreOffice нет (локальная разработка), возвращаем пустой PDF
        # или можно генерировать простой PDF через reportlab
        return None
